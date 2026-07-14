from docx import Document
import os
from typing import Dict, Any

class WordParser:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract all text from Word document"""
        text = ""
        try:
            file_path = os.path.normpath(file_path)
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Word extraction failed: {str(e)}")
        return text
    
    @staticmethod
    def extract_tables(file_path: str) -> list:
        """Extract tables from Word document"""
        tables = []
        try:
            file_path = os.path.normpath(file_path)
            doc = Document(file_path)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
        except Exception as e:
            print(f"Word table extraction warning: {str(e)}")
        return tables

    @staticmethod
    def extract_full(file_path: str) -> Dict[str, any]:
        """Complete extraction pipeline"""
        file_path = os.path.normpath(file_path)
        text = WordParser.extract_text(file_path)
        
        if not text or len(text.strip()) < 50:
            raise ValueError("Could not extract sufficient text from Word document")
        
        from app.parsers.pdf_parser import PDFParser
        structured = PDFParser.parse_syllabus_structure(text)
        tables = WordParser.extract_tables(file_path)
        
        return {
            "raw_text": text,
            "structured": structured,
            "tables": tables,
            "word_count": len(text.split())
        }