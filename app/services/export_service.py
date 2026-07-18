import os
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

from app.config import settings

class ExportService:
    def __init__(self):
        self.output_dir = settings.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)
        self.styles = getSampleStyleSheet()

    def generate_json(self, cdp) -> str:
        """Generate a JSON file from CDP data."""
        filename = f"CDP_{cdp.course_code}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.json"
        output_path = os.path.join(self.output_dir, filename)

        export_data = {
            "id": cdp.id,
            "syllabus_id": cdp.syllabus_id,
            "course_name": cdp.course_name,
            "course_code": cdp.course_code,
            "credits": cdp.credits,
            "status": cdp.status,
            "created_at": cdp.created_at.isoformat() if cdp.created_at else None,
            "updated_at": cdp.updated_at.isoformat() if cdp.updated_at else None,
            "cdp_json": cdp.cdp_json,
            "concept_map": cdp.concept_map,
        }

        with open(output_path, "w", encoding="utf-8") as file:
            json.dump(export_data, file, indent=2, ensure_ascii=False)

        return output_path
    
    # ============ PDF GENERATION (COMPLETE) ============
    def generate_pdf(self, cdp) -> str:
        """Generate professional PDF from CDP data"""
        filename = f"CDP_{cdp.course_code}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4, 
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
        story = []
        
        cdp_data = cdp.cdp_json
        
        # Add all sections
        story.extend(self._create_header())
        story.append(Spacer(1, 0.3*inch))
        story.extend(self._create_title())
        story.append(Spacer(1, 0.2*inch))
        story.extend(self._create_course_info(cdp_data))
        story.append(Spacer(1, 0.2*inch))
        story.extend(self._create_co_section(cdp_data))
        story.append(Spacer(1, 0.2*inch))
        story.extend(self._create_po_section(cdp_data))
        story.append(Spacer(1, 0.2*inch))
        story.extend(self._create_weekly_plan(cdp_data))
        story.append(PageBreak())
        story.extend(self._create_affinity_map(cdp_data))
        story.append(Spacer(1, 0.2*inch))
        story.extend(self._create_evaluation_scheme(cdp_data))
        story.append(Spacer(1, 0.2*inch))
        story.extend(self._create_concept_map_section(cdp_data))
        
        doc.build(story)
        return filepath
    
    def _create_header(self):
        style = ParagraphStyle(
            'Header',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            alignment=1
        )
        return [Paragraph("AcadPlan AI - Automated Course Delivery Plan Generator", style)]
    
    def _create_title(self):
        style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.darkblue,
            spaceAfter=20,
            alignment=1
        )
        return [Paragraph("COURSE DELIVERY PLAN", style)]
    
    def _create_course_info(self, data):
        style = self.styles['Normal']
        info = [
            f"<b>Course Name:</b> {data.get('course_name', 'N/A')}",
            f"<b>Course Code:</b> {data.get('course_code', 'N/A')}",
            f"<b>Credits (L-T-P-C):</b> {data.get('credits', 'N/A')}",
            f"<b>Department:</b> {data.get('department', 'N/A')}",
            f"<b>Academic Year:</b> {data.get('academic_year', 'N/A')}",
            f"<b>Prerequisites:</b> {', '.join(data.get('prerequisites', ['None']))}"
        ]
        return [Paragraph(line, style) for line in info]
    
    def _create_co_section(self, data):
        style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=10
        )
        content = [Paragraph("1. COURSE OUTCOMES (COs)", style)]
        cos = data.get('course_outcomes', [])
        for co in cos:
            pos = ', '.join(co.get('mapped_pos', []))
            content.append(Paragraph(
                f"- <b>{co.get('id', '')}:</b> {co.get('description', '')} "
                f"<i>(Mapped to: {pos})</i>",
                self.styles['Normal']
            ))
        return content
    
    def _create_po_section(self, data):
        style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=10
        )
        content = [Paragraph("2. PROGRAM OUTCOMES (POs)", style)]
        pos = data.get('program_outcomes', [])
        for po in pos:
            content.append(Paragraph(
                f"- <b>{po.get('id', '')}:</b> {po.get('description', '')}",
                self.styles['Normal']
            ))
        return content
    
    def _create_weekly_plan(self, data):
        style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=10
        )
        content = [Paragraph("3. WEEKLY DELIVERY PLAN", style)]
        
        weekly_plan = data.get('weekly_plan', [])
        for week in weekly_plan:
            content.append(Paragraph(
                f"<b>Week {week.get('week', '')}</b> - Total Hours: {week.get('total_hours', 'N/A')}",
                self.styles['Heading3']
            ))
            
            topics = week.get('topics', [])
            if topics:
                table_data = [['Topic', 'Duration', 'Objectives']]
                for topic in topics:
                    objectives = ', '.join(topic.get('learning_objectives', [])[:2])
                    table_data.append([
                        f"{topic.get('id', '')}: {topic.get('title', '')}",
                        topic.get('duration', 'N/A'),
                        objectives
                    ])
                
                table = Table(table_data, colWidths=[2.5*inch, 1*inch, 2.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('TOPPADDING', (0, 0), (-1, -1), 5),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 5)
                ]))
                content.append(table)
            
            content.append(Spacer(1, 0.1*inch))
        
        return content
    
    def _create_affinity_map(self, data):
        style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=10
        )
        content = [Paragraph("4. CO-PO AFFINITY MAP", style)]
        
        affinity = data.get('co_po_affinity_map', {})
        if affinity:
            import re
            all_pos = set()
            for po_scores in affinity.values():
                all_pos.update(po_scores.keys())
            pos_list = sorted(list(all_pos), key=lambda x: (0 if x.startswith('PO') else 1, int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else x))
            table_data = [['COs \\ POs'] + list(pos_list)]
            for co_id, po_scores in affinity.items():
                row = [co_id]
                for po in pos_list:
                    score = po_scores.get(po, 0)
                    row.append(f"{score:.2f}")
                table_data.append(row)
            
            table = Table(table_data, colWidths=[0.8*inch] + [0.7*inch] * len(pos_list))
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9)
            ]))
            content.append(table)
        
        return content
    
    def _create_evaluation_scheme(self, data):
        style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=10
        )
        content = [Paragraph("5. EVALUATION SCHEME", style)]
        
        scheme = data.get('evaluation_scheme', [])
        if scheme:
            table_data = [['Component', 'Weightage', 'Duration']]
            for item in scheme:
                table_data.append([
                    item.get('component', 'N/A'),
                    f"{item.get('weightage', 0)}%",
                    item.get('duration', 'N/A') or 'N/A'
                ])
            
            table = Table(table_data, colWidths=[2*inch, 1.5*inch, 2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey)
            ]))
            content.append(table)
        
        return content
    
    def _create_concept_map_section(self, data):
        style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=10
        )
        content = [Paragraph("6. CONCEPT MAP", style)]
        content.append(Paragraph("Concept map visualization available in web interface.", self.styles['Normal']))
        content.append(Paragraph("Mermaid.js Code:", self.styles['Normal']))
        
        code_style = ParagraphStyle(
            'Code',
            parent=self.styles['Normal'],
            fontSize=8,
            fontName='Courier',
            backColor=colors.lightgrey,
            leftIndent=20,
            rightIndent=20,
            spaceAfter=10
        )
        
        concept_map = data.get('concept_map_mermaid', '')
        if concept_map:
            for line in concept_map.strip().split('\n'):
                if line.strip():
                    content.append(Paragraph(line.strip(), code_style))
        
        return content

    # ============ WORD GENERATION (COMPLETE) ============
    def generate_docx(self, cdp) -> str:
        """Generate professional Word document from CDP data"""
        filename = f"CDP_{cdp.course_code}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = DocxDocument()
        cdp_data = cdp.cdp_json
        
        # === Title ===
        title = doc.add_heading('COURSE DELIVERY PLAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        # === Course Info ===
        doc.add_heading('Course Information', level=1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ('Course Name', cdp_data.get('course_name', 'N/A')),
            ('Course Code', cdp_data.get('course_code', 'N/A')),
            ('Credits (L-T-P-C)', cdp_data.get('credits', 'N/A')),
            ('Department', cdp_data.get('department', 'N/A')),
            ('Academic Year', cdp_data.get('academic_year', 'N/A')),
            ('Prerequisites', ', '.join(cdp_data.get('prerequisites', ['None'])))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # === Course Outcomes ===
        doc.add_heading('1. COURSE OUTCOMES (COs)', level=1)
        cos = cdp_data.get('course_outcomes', [])
        for co in cos:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{co.get('id', '')}: ").bold = True
            p.add_run(f"{co.get('description', '')} ")
            pos = ', '.join(co.get('mapped_pos', []))
            p.add_run(f"(Mapped to: {pos})").italic = True
        
        doc.add_paragraph()
        
        # === Program Outcomes ===
        doc.add_heading('2. PROGRAM OUTCOMES (POs)', level=1)
        pos = cdp_data.get('program_outcomes', [])
        for po in pos:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{po.get('id', '')}: ").bold = True
            p.add_run(po.get('description', ''))
        
        doc.add_paragraph()
        
        # === Weekly Plan ===
        doc.add_heading('3. WEEKLY DELIVERY PLAN', level=1)
        weekly_plan = cdp_data.get('weekly_plan', [])
        for week in weekly_plan:
            doc.add_heading(f"Week {week.get('week', '')} - Total Hours: {week.get('total_hours', 'N/A')}", level=2)
            
            topics = week.get('topics', [])
            if topics:
                table = doc.add_table(rows=len(topics)+1, cols=3)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header
                headers = ['Topic', 'Duration', 'Objectives']
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                
                # Data
                for row_idx, topic in enumerate(topics, 1):
                    table.cell(row_idx, 0).text = f"{topic.get('id', '')}: {topic.get('title', '')}"
                    table.cell(row_idx, 1).text = topic.get('duration', 'N/A')
                    objectives = ', '.join(topic.get('learning_objectives', [])[:2])
                    table.cell(row_idx, 2).text = objectives
            
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # === CO-PO Affinity Map ===
        doc.add_heading('4. CO-PO AFFINITY MAP', level=1)
        affinity = cdp_data.get('co_po_affinity_map', {})
        if affinity:
            import re
            all_pos = set()
            for po_scores in affinity.values():
                all_pos.update(po_scores.keys())
            pos_list = sorted(list(all_pos), key=lambda x: (0 if x.startswith('PO') else 1, int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else x))            
            table = doc.add_table(rows=len(affinity)+1, cols=len(pos_list)+1)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header
            table.cell(0, 0).text = 'COs \\ POs'
            for i, po in enumerate(pos_list, 1):
                table.cell(0, i).text = po
                table.cell(0, i).paragraphs[0].runs[0].bold = True
            # Data
            for row_idx, (co_id, po_scores) in enumerate(affinity.items(), 1):
                table.cell(row_idx, 0).text = co_id
                table.cell(row_idx, 0).paragraphs[0].runs[0].bold = True
                for col_idx, po in enumerate(pos_list, 1):
                    score = po_scores.get(po, 0)
                    table.cell(row_idx, col_idx).text = f"{score:.2f}"
                    
                    # Color code based on score
                    cell = table.cell(row_idx, col_idx)
                    if score >= 0.8:
                        self._set_cell_shading(cell, '00CC00')  # Green
                    elif score >= 0.5:
                        self._set_cell_shading(cell, 'FFD700')  # Gold
                    else:
                        self._set_cell_shading(cell, 'FF6B6B')  # Red
        
        doc.add_paragraph()
        
        # === Evaluation Scheme ===
        doc.add_heading('5. EVALUATION SCHEME', level=1)
        scheme = cdp_data.get('evaluation_scheme', [])
        if scheme:
            table = doc.add_table(rows=len(scheme)+1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header
            headers = ['Component', 'Weightage', 'Duration']
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Data
            for row_idx, item in enumerate(scheme, 1):
                table.cell(row_idx, 0).text = item.get('component', 'N/A')
                table.cell(row_idx, 1).text = f"{item.get('weightage', 0)}%"
                table.cell(row_idx, 2).text = item.get('duration', 'N/A') or 'N/A'
        
        doc.add_paragraph()
        
        # === Concept Map ===
        doc.add_heading('6. CONCEPT MAP', level=1)
        doc.add_paragraph("Concept map visualization available in web interface.")
        doc.add_paragraph("Mermaid.js Code:", style='Normal')
        doc.add_run().bold = True
        
        concept_map = cdp_data.get('concept_map_mermaid', '')
        if concept_map:
            p = doc.add_paragraph()
            for line in concept_map.strip().split('\n'):
                if line.strip():
                    run = p.add_run(line.strip() + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)
        
        # Save document
        doc.save(filepath)
        return filepath
    
    def _set_cell_shading(self, cell, color_hex):
        """Set background color for Word table cell"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading_elm)

